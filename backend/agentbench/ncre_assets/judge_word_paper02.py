# NCRE Word 操作题私有判分脚本（经典题库第2套：领慧讲堂海报，command_metrics 协议）
# 检查 WORD.docx 的自定义纸张、页边距、背景填充、标题格式、日程表与流程列表；
# 产物缺失或损坏时对应指标为 0。
import contextlib
import json
import re
import sys
import zipfile

with contextlib.suppress(Exception):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

metrics = {key: 0.0 for key in ["w1", "w2", "w3", "w4", "w5", "w6", "w7", "w8", "w9"]}
evidence = {}

CM_TOL = 0.1
PT_TOL = 1.0
SCHEDULE_TOPICS = ["签到", "大学生职场定位和职业准备", "大学生人生规划", "现场提问"]
SIGNUP_STEPS = ["学工处报名", "确认坐席", "领取资料", "领取门票"]


def emit():
    print("AGENTBENCH_METRICS=" + json.dumps({"metrics": metrics, "evidence": evidence}))


def approx(actual, expected, tol):
    try:
        return abs(float(actual) - float(expected)) <= tol
    except (TypeError, ValueError):
        return False


try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except Exception as exc:  # noqa: BLE001
    evidence["error"] = f"python-docx 不可用: {str(exc)[:200]}"
    emit()
    raise SystemExit(0) from None


def east_asia_font(run):
    rpr = run._element.rPr
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:eastAsia"))


document = None
try:
    document = Document("WORD.docx")
except Exception as exc:  # noqa: BLE001
    evidence["error"] = f"WORD.docx 缺失或无法打开: {str(exc)[:200]}"
    emit()
    raise SystemExit(0) from None

# ------------------------------------------------ w1 第1页自定义纸张 27x35cm
try:
    section = document.sections[0]
    if approx(section.page_width.cm, 27, CM_TOL) and approx(section.page_height.cm, 35, CM_TOL):
        metrics["w1"] = 100.0
    evidence["w1"] = f"{section.page_width.cm:.2f}cm x {section.page_height.cm:.2f}cm"
except Exception as exc:  # noqa: BLE001
    evidence["w1"] = str(exc)[:200]

# ------------------------------------------------ w2 第1页页边距 上下5 左右3
try:
    section = document.sections[0]
    checks = [
        approx(section.top_margin.cm, 5, CM_TOL),
        approx(section.bottom_margin.cm, 5, CM_TOL),
        approx(section.left_margin.cm, 3, CM_TOL),
        approx(section.right_margin.cm, 3, CM_TOL),
    ]
    metrics["w2"] = 100.0 * sum(checks) / len(checks)
    evidence["w2"] = f"{sum(checks)}/4 边距正确"
except Exception as exc:  # noqa: BLE001
    evidence["w2"] = str(exc)[:200]

# ------------------------------------------------ w3 页面背景纯色填充 FFF2CC
try:
    with zipfile.ZipFile("WORD.docx") as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
    background = re.search(r"<w:background[^>]*>", document_xml)
    if background and "fff2cc" in background.group(0).lower():
        metrics["w3"] = 100.0
    evidence["w3"] = (background.group(0)[:120] if background else "无 w:background")
except Exception as exc:  # noqa: BLE001
    evidence["w3"] = str(exc)[:200]

paragraphs = document.paragraphs

# ------------------------------------------------ w4 主标题 微软雅黑62磅红色
try:
    title = next((p for p in paragraphs if p.text.strip() == "领慧讲堂"), None)
    checks = []
    if title is not None and title.runs:
        fonts = {run.font.name for run in title.runs} | {east_asia_font(run) for run in title.runs}
        checks.append("微软雅黑" in fonts)
        checks.append(any(
            run.font.size is not None and approx(run.font.size.pt, 62, PT_TOL)
            for run in title.runs))
        checks.append(any(
            run.font.color is not None and run.font.color.rgb is not None
            and str(run.font.color.rgb).upper() == "FF0000" for run in title.runs))
    metrics["w4"] = round(100.0 * sum(checks) / 3, 2) if checks else 0.0
    evidence["w4"] = f"主标题子项 {sum(checks)}/3" if checks else "未找到主标题段"
except Exception as exc:  # noqa: BLE001
    evidence["w4"] = str(exc)[:200]

# ------------------------------------------------ w5 第2页 A4 横向
try:
    if len(document.sections) >= 2:
        page2 = document.sections[1]
        checks = [
            page2.orientation == WD_ORIENT.LANDSCAPE,
            approx(page2.page_width.cm, 29.7, CM_TOL),
            approx(page2.page_height.cm, 21, CM_TOL),
        ]
        metrics["w5"] = 100.0 * sum(checks) / len(checks)
        evidence["w5"] = f"{page2.page_width.cm:.2f}x{page2.page_height.cm:.2f} orient={page2.orientation}"
    else:
        evidence["w5"] = f"仅 {len(document.sections)} 节"
except Exception as exc:  # noqa: BLE001
    evidence["w5"] = str(exc)[:200]

# ------------------------------------------------ w6 报告人姓名 赵蕈
try:
    metrics["w6"] = 100.0 if any("赵蕈" in p.text for p in paragraphs) else 0.0
    evidence["w6"] = "含赵蕈" if metrics["w6"] else "未找到赵蕈"
except Exception as exc:  # noqa: BLE001
    evidence["w6"] = str(exc)[:200]

# ------------------------------------------------ w7 日程安排表 4行3列
try:
    for table in document.tables:
        if len(table.rows) != 5 or len(table.columns) != 3:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        checks = [header == ["时间", "主题", "报告人"]]
        topics = [table.rows[i].cells[1].text.strip() for i in range(1, 5)]
        checks.append(topics == SCHEDULE_TOPICS)
        metrics["w7"] = 100.0 * sum(checks) / len(checks)
        evidence["w7"] = f"表头{header} 主题{topics}"
        break
    else:
        evidence["w7"] = "未找到4行3列表格"
except Exception as exc:  # noqa: BLE001
    evidence["w7"] = str(exc)[:200]

# ------------------------------------------------ w8 报名流程编号列表
try:
    texts = [p.text for p in paragraphs]
    hits = sum(1 for step in SIGNUP_STEPS if any(step in text for text in texts))
    metrics["w8"] = 100.0 * hits / len(SIGNUP_STEPS)
    evidence["w8"] = f"{hits}/4 流程项"
except Exception as exc:  # noqa: BLE001
    evidence["w8"] = str(exc)[:200]

# ------------------------------------------------ w9 报告人介绍 两端对齐+首行缩进2字符
try:
    intro = next((p for p in paragraphs if "资深媒体人" in p.text), None)
    if intro is not None:
        align_ok = intro.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        indent = intro.paragraph_format.first_line_indent
        indent_ok = (indent is not None and approx(indent.pt, 21, PT_TOL)) \
            or 'firstLineChars="200"' in intro._p.xml
        metrics["w9"] = 50.0 * align_ok + 50.0 * indent_ok
        evidence["w9"] = f"对齐={intro.alignment}, 缩进ok={indent_ok}"
    else:
        evidence["w9"] = "未找到报告人介绍段"
except Exception as exc:  # noqa: BLE001
    evidence["w9"] = str(exc)[:200]

emit()
