# NCRE Word 操作题私有判分脚本（经典题库第4套：统计工作年报排版，command_metrics 协议）
# 检查 年报.docx 的空格删除、16开页面、封面分页、咨询表+饼图、标题样式、
# 超链接+脚注、两栏、目录域与奇偶页眉；产物缺失时全 0。
import contextlib
import json
import re
import sys
import zipfile

with contextlib.suppress(Exception):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

metrics = {key: 0.0 for key in ["w1", "w2", "w3", "w4", "w5", "w6", "w7", "w8", "w9"]}
evidence = {}

CM_TOL = 0.1
LINK_TEXT = "统计局队政府网站"
LINK_TARGET = "http://www.bjstats.gov.cn/"
COVER_MARK = "二〇一三年三月"
CONSULT_HEADER = ["咨询方式", "人次", "所占比例(%)"]
CONSULT_DATA = [("现场咨询", "93", "5.04"), ("电话咨询", "1515", "82.07"),
                ("网上咨询", "238", "12.89"), ("合计", "1846", "100")]
HEADER_TEXT = "北京市政府信息公开工作年度报告"


def emit():
    print("AGENTBENCH_METRICS=" + json.dumps({"metrics": metrics, "evidence": evidence}))


def approx(actual, expected, tol):
    try:
        return abs(float(actual) - float(expected)) <= tol
    except (TypeError, ValueError):
        return False


def cell_text(cell):
    return cell.text.strip()


def close_text(actual, expected):
    try:
        return abs(float(actual) - float(expected)) < 0.005
    except (TypeError, ValueError):
        return str(actual).strip() == str(expected).strip()


try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as exc:  # noqa: BLE001
    evidence["error"] = f"python-docx 不可用: {str(exc)[:200]}"
    emit()
    raise SystemExit(0) from None

document = None
try:
    document = Document("年报.docx")
except Exception as exc:  # noqa: BLE001
    evidence["error"] = f"年报.docx 缺失或无法打开: {str(exc)[:200]}"
    emit()
    raise SystemExit(0) from None

paragraphs = document.paragraphs

# ------------------------------------------------ w1 删除汉字与西文/数字间半角空格
try:
    non_empty = [p for p in paragraphs if p.text.strip()]
    clean = [p for p in non_empty if " " not in p.text]
    metrics["w1"] = round(100.0 * len(clean) / len(non_empty), 2) if non_empty else 0.0
    evidence["w1"] = f"{len(clean)}/{len(non_empty)} 段无半角空格"
except Exception as exc:  # noqa: BLE001
    evidence["w1"] = str(exc)[:200]

# ------------------------------------------------ w2 16开纸张与页边距
try:
    section = document.sections[0]
    checks = [
        approx(section.page_width.cm, 18.4, CM_TOL),
        approx(section.page_height.cm, 26, CM_TOL),
        approx(section.top_margin.cm, 3.2, CM_TOL),
        approx(section.bottom_margin.cm, 3, CM_TOL),
        approx(section.left_margin.cm, 2.5, CM_TOL),
        approx(section.right_margin.cm, 2.5, CM_TOL),
    ]
    metrics["w2"] = round(100.0 * sum(checks) / len(checks), 2)
    evidence["w2"] = f"{sum(checks)}/6 项正确"
except Exception as exc:  # noqa: BLE001
    evidence["w2"] = str(exc)[:200]

# ------------------------------------------------ w3 封面独占一页
try:
    cover_index = next(
        (i for i, p in enumerate(paragraphs) if COVER_MARK in p.text), None)
    if cover_index is None:
        evidence["w3"] = "未找到封面段"
    else:
        break_in_cover = any(
            'w:type="page"' in paragraphs[i]._p.xml for i in range(cover_index + 1))
        next_break = (cover_index + 1 < len(paragraphs)
                      and paragraphs[cover_index + 1].paragraph_format.page_break_before)
        metrics["w3"] = 100.0 if (break_in_cover or next_break) else 0.0
        evidence["w3"] = f"封底段内分页符={break_in_cover}, 后段pageBreakBefore={next_break}"
except Exception as exc:  # noqa: BLE001
    evidence["w3"] = str(exc)[:200]

# ------------------------------------------------ w4 咨询情况表与饼图
try:
    table_ok = data_ok = False
    for table in document.tables:
        if len(table.rows) != 5 or len(table.columns) != 3:
            continue
        header = [cell_text(cell) for cell in table.rows[0].cells]
        if header != CONSULT_HEADER:
            continue
        table_ok = True
        rows_ok = []
        for i, expected in enumerate(CONSULT_DATA, start=1):
            cells = [cell_text(cell) for cell in table.rows[i].cells]
            rows_ok.append(
                cells[0] == expected[0] and close_text(cells[1], expected[1])
                and close_text(cells[2], expected[2]))
        data_ok = all(rows_ok)
        evidence["w4"] = f"表头ok={table_ok}, 数据行{sum(rows_ok)}/4"
        break
    if not table_ok:
        evidence["w4"] = "未找到咨询情况表"
    chart_ok = False
    with zipfile.ZipFile("年报.docx") as archive:
        for name in archive.namelist():
            if not re.search(r"word/charts/chart\d+\.xml$", name):
                continue
            chart_xml = archive.read(name).decode("utf-8", errors="replace")
            if "pieChart" in chart_xml \
                    and re.search(r'<(?:c:)?showPercent val="1"', chart_xml) \
                    and re.search(r'<(?:c:)?showVal val="0"', chart_xml):
                chart_ok = True
                break
    metrics["w4"] = round(40.0 * table_ok + 30.0 * data_ok + 30.0 * chart_ok, 2)
    evidence["w4"] = evidence.get("w4", "") + f", 饼图ok={chart_ok}"
except Exception as exc:  # noqa: BLE001
    evidence["w4"] = str(exc)[:200]

# ------------------------------------------------ w5 标题样式层级
try:
    total = matched = 0
    for paragraph in paragraphs:
        text = paragraph.text.strip()
        if re.match(r"^[一二三四五六七八九十]+、", text):
            level = 1
        elif text.startswith("（"):
            level = 2
        elif re.match(r"^\d+、", text):
            level = 3
        else:
            continue
        total += 1
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name in (f"Heading {level}", f"标题 {level}"):
            matched += 1
    metrics["w5"] = round(100.0 * matched / total, 2) if total else 0.0
    evidence["w5"] = f"{matched}/{total} 标题样式正确"
except Exception as exc:  # noqa: BLE001
    evidence["w5"] = str(exc)[:200]

# ------------------------------------------------ w6 超链接与脚注
try:
    with zipfile.ZipFile("年报.docx") as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        rels_xml = archive.read(
            "word/_rels/document.xml.rels").decode("utf-8", errors="replace")
        footnotes_xml = ""
        if "word/footnotes.xml" in archive.namelist():
            footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8", errors="replace")
    targets = dict(re.findall(r'Id="([^"]+)"[^>]*Target="([^"]+)"', rels_xml))
    link_ok = text_ok = color_ok = False
    for rid, inner in re.findall(
            r'<w:hyperlink[^>]*r:id="([^"]+)"[^>]*>(.*?)</w:hyperlink>',
            document_xml, re.DOTALL):
        if targets.get(rid) != LINK_TARGET:
            continue
        link_ok = True
        plain = re.sub(r"<[^>]+>", "", inner)
        if LINK_TEXT in plain:
            text_ok = True
        if "FF0000" in inner.upper():
            color_ok = True
    note_ok = False
    for block in re.findall(r"<w:footnote\b[^>]*>(.*?)</w:footnote>", footnotes_xml, re.DOTALL):
        if 'type="separator"' in block or 'type="continuationSeparator"' in block:
            continue
        note_text = "".join(re.findall(r"<w:t[^>]*>([^<]*)</w:t>", block))
        if note_text.strip():
            note_ok = True
            break
    metrics["w6"] = 25.0 * sum([link_ok, text_ok, color_ok, note_ok])
    evidence["w6"] = f"链接={link_ok}, 文本={text_ok}, 红色={color_ok}, 脚注={note_ok}"
except Exception as exc:  # noqa: BLE001
    evidence["w6"] = str(exc)[:200]

# ------------------------------------------------ w7 正文两栏
try:
    with zipfile.ZipFile("年报.docx") as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
    two_cols = re.search(r'<w:cols[^>]*w:num="2"', document_xml) is not None
    metrics["w7"] = 100.0 if two_cols else 0.0
    evidence["w7"] = f"w:cols num=2: {two_cols}"
except Exception as exc:  # noqa: BLE001
    evidence["w7"] = str(exc)[:200]

# ------------------------------------------------ w8 目录域
try:
    toc = re.search(r"<w:instrText[^>]*>[^<]*TOC", document_xml) is not None
    metrics["w8"] = 100.0 if toc else 0.0
    evidence["w8"] = f"TOC域: {toc}"
except Exception as exc:  # noqa: BLE001
    evidence["w8"] = str(exc)[:200]

# ------------------------------------------------ w9 页眉与页码（奇右偶左）
try:
    with zipfile.ZipFile("年报.docx") as archive:
        settings_xml = archive.read("word/settings.xml").decode("utf-8", errors="replace")
    even_odd = "<w:evenAndOddHeaders" in settings_xml

    def header_check(header, expected_align):
        for paragraph in header.paragraphs:
            if "PAGE" in paragraph._p.xml and paragraph.alignment == expected_align:
                return True
        return False

    section = document.sections[-1]
    text_ok = any(
        HEADER_TEXT in paragraph.text
        for header in (section.header, section.even_page_header)
        for paragraph in header.paragraphs)
    odd_ok = header_check(section.header, WD_ALIGN_PARAGRAPH.RIGHT)
    even_ok = header_check(section.even_page_header, WD_ALIGN_PARAGRAPH.LEFT)
    metrics["w9"] = 25.0 * sum([even_odd, text_ok, odd_ok, even_ok])
    evidence["w9"] = f"奇偶页眉={even_odd}, 文字={text_ok}, 奇右={odd_ok}, 偶左={even_ok}"
except Exception as exc:  # noqa: BLE001
    evidence["w9"] = str(exc)[:200]

emit()
