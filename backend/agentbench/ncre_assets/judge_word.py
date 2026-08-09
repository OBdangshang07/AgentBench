# NCRE Word 操作题私有判分脚本（AgentBench command_metrics 协议）
# 检查 Word.docx（w1-w8）与 Word-邀请函.docx（w9）；产物缺失时对应指标为 0。
import contextlib
import json
import re
import sys
import zipfile

with contextlib.suppress(Exception):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

metrics = {key: 0.0 for key in ["w1", "w2", "w3", "w4", "w5", "w6", "w7", "w8", "w9"]}
evidence = {}

CONTACT_NAMES = ["李达志", "王建国", "刘晓梅", "陈志强", "赵雅琴"]
CM_TOL = 0.1
PT_TOL = 1.0


def emit():
    print("AGENTBENCH_METRICS=" + json.dumps({"metrics": metrics, "evidence": evidence}))


def approx(actual, expected, tol):
    try:
        return abs(float(actual) - float(expected)) <= tol
    except (TypeError, ValueError):
        return False


try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except Exception as exc:  # noqa: BLE001
    evidence["error"] = f"python-docx 不可用: {str(exc)[:200]}"
    emit()
    raise SystemExit(0) from None


def east_asia_font(run):
    rpr = run._element.rPr
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:eastAsia"))


# ------------------------------------------------ Word.docx 模板检查
template = None
try:
    template = Document("Word.docx")
except Exception as exc:  # noqa: BLE001
    evidence["Word.docx"] = f"无法打开: {str(exc)[:200]}"

if template is not None:
    section = template.sections[0]
    # w1 页面尺寸
    if approx(section.page_width.cm, 30, CM_TOL) and approx(section.page_height.cm, 18, CM_TOL):
        metrics["w1"] = 100.0
    evidence["w1"] = f"{section.page_width.cm:.2f}cm x {section.page_height.cm:.2f}cm"
    # w2 页边距
    checks = [
        approx(section.top_margin.cm, 2, CM_TOL),
        approx(section.bottom_margin.cm, 2, CM_TOL),
        approx(section.left_margin.cm, 3, CM_TOL),
        approx(section.right_margin.cm, 3, CM_TOL),
    ]
    metrics["w2"] = 100.0 * sum(checks) / len(checks)
    # w3 页面背景填充 FDE9D9
    try:
        with zipfile.ZipFile("Word.docx") as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        background = re.search(r"<w:background[^>]*>", document_xml)
        if background and "fde9d9" in background.group(0).lower():
            metrics["w3"] = 100.0
        evidence["w3"] = (background.group(0)[:120] if background else "无 w:background")
    except Exception as exc:  # noqa: BLE001
        evidence["w3"] = str(exc)[:200]

    paragraphs = template.paragraphs
    # w4 全文字体微软雅黑
    total_runs = 0
    ok_runs = 0
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            total_runs += 1
            if run.font.name == "微软雅黑" or east_asia_font(run) == "微软雅黑":
                ok_runs += 1
    if total_runs:
        metrics["w4"] = round(100.0 * ok_runs / total_runs, 2)
    evidence["w4"] = f"{ok_runs}/{total_runs} runs"
    # w5 标题字号一号(26磅) + 第1段蓝色
    checks = []
    for index in (0, 1):
        size_ok = False
        if index < len(paragraphs):
            for run in paragraphs[index].runs:
                if run.font.size is not None and approx(run.font.size.pt, 26, 0.5):
                    size_ok = True
        checks.append(size_ok)
    blue_ok = False
    if paragraphs:
        for run in paragraphs[0].runs:
            color = run.font.color
            if color is not None and (
                (color.rgb is not None and str(color.rgb) == "0000FF")
                or color.theme_color is not None
            ):
                blue_ok = True
    checks.append(blue_ok)
    metrics["w5"] = round(100.0 * sum(checks) / len(checks), 2)
    # w6 对齐方式：标题居中，落款与日期右对齐
    checks = []
    for index, expected in ((0, WD_ALIGN_PARAGRAPH.CENTER), (1, WD_ALIGN_PARAGRAPH.CENTER),
                            (5, WD_ALIGN_PARAGRAPH.RIGHT), (6, WD_ALIGN_PARAGRAPH.RIGHT)):
        checks.append(
            index < len(paragraphs) and paragraphs[index].alignment == expected
        )
    metrics["w6"] = 100.0 * sum(checks) / len(checks)
    # w7 正文段（第4、5段）首行缩进2字符
    checks = []
    for index in (3, 4):
        ok = False
        if index < len(paragraphs):
            paragraph = paragraphs[index]
            indent = paragraph.paragraph_format.first_line_indent
            if (indent is not None and approx(indent.pt, 21, PT_TOL)) or \
                    'firstLineChars="200"' in paragraph._p.xml:
                ok = True
        checks.append(ok)
    metrics["w7"] = 100.0 * sum(checks) / len(checks)
    # w8 标题段前段后各0.5行
    checks = []
    for index in (0, 1):
        if index >= len(paragraphs):
            checks.extend([False, False])
            continue
        paragraph = paragraphs[index]
        xml_text = paragraph._p.xml
        fmt = paragraph.paragraph_format
        before_ok = (fmt.space_before is not None and approx(fmt.space_before.pt, 13, PT_TOL)) \
            or "beforeLines=\"50\"" in xml_text
        after_ok = (fmt.space_after is not None and approx(fmt.space_after.pt, 13, PT_TOL)) \
            or "afterLines=\"50\"" in xml_text
        checks.extend([before_ok, after_ok])
    metrics["w8"] = 100.0 * sum(checks) / len(checks)

# ------------------------------------------------ Word-邀请函.docx 合并结果检查
try:
    merged = Document("Word-邀请函.docx")
    with zipfile.ZipFile("Word-邀请函.docx") as archive:
        merged_xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
    salutations = [p.text for p in merged.paragraphs if "尊敬的" in p.text]
    name_score = 0.0
    for name in CONTACT_NAMES:
        hits = [text for text in salutations if name in text]
        if len(hits) == 1 and not any(
            other != name and other in hits[0] for other in CONTACT_NAMES
        ):
            name_score += 1.0
    name_ratio = name_score / len(CONTACT_NAMES)
    count_ok = len(salutations) == len(CONTACT_NAMES)
    page_breaks = merged_xml.count('w:type="page"')
    page_ok = page_breaks >= len(CONTACT_NAMES) - 1
    metrics["w9"] = round(60.0 * name_ratio + 20.0 * count_ok + 20.0 * page_ok, 2)
    evidence["w9"] = f"称呼段{len(salutations)}个, 分页符{page_breaks}个, 姓名匹配{name_score:.0f}/{len(CONTACT_NAMES)}"
except Exception as exc:  # noqa: BLE001
    evidence["w9"] = f"Word-邀请函.docx 缺失或无法打开: {str(exc)[:200]}"

emit()
